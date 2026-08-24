# -*- coding: utf-8 -*-
import json
import os
import re
import sys
from copy import deepcopy

import docx
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
import content_blocks as C

TEMPLATE = os.path.join(HERE, "AIA_template.docx")
FIGDIR = os.path.join(HERE, "figures")
OUT = os.path.join(HERE, "AIA_MPIR_manuscript.docx")
CITMAP = json.load(open(os.path.join(HERE, "citation_map.json"), encoding="utf-8"))

CITE_RE = re.compile(r"‹([\d,]+)›")  # ‹...›

# ---------------------------------------------------------------------------
# Citation numbering: first-appearance order across the whole BLOCKS list.
# ---------------------------------------------------------------------------
order = []
order_set = set()


def scan_text_for_citations(text):
    for m in CITE_RE.finditer(text):
        for key in m.group(1).split(","):
            k = f"reference_{key.strip()}"
            if k not in order_set:
                order_set.add(k)
                order.append(k)


def all_text_in_blocks():
    for b in C.BLOCKS:
        if b["type"] in ("p",):
            yield b["text"]
        elif b["type"] == "bullets":
            for head, body in b["items"]:
                yield head + " " + body
        elif b["type"] == "figure":
            yield b["caption"]
        elif b["type"] == "table":
            yield b["caption"]
        elif b["type"] == "code":
            if b.get("caption"):
                yield b["caption"]


for t in all_text_in_blocks():
    scan_text_for_citations(t)

key_to_num = {k: i + 1 for i, k in enumerate(order)}
print(f"Resolved {len(order)} unique citations in first-appearance order.", file=sys.stderr)


def resolve_citations(text):
    def _sub(m):
        nums = [str(key_to_num[f"reference_{key.strip()}"]) for key in m.group(1).split(",")]
        return "[" + ", ".join(nums) + "]"
    return CITE_RE.sub(_sub, text)


# ---------------------------------------------------------------------------
doc = docx.Document(TEMPLATE)

# ---------------------------------------------------------------------------
# 1. Front matter edits (paragraphs 0-24 kept; edit text in place)
# ---------------------------------------------------------------------------
paras = doc.paragraphs


def set_para_text(p, text):
    # Keep the formatting of the first run, drop the rest, set new text.
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""


set_para_text(paras[0], "Received: 22 August 2026 | Revised: TBD | Accepted: TBD | Published online: TBD")
set_para_text(paras[3], "RESEARCH ARTICLE")
set_para_text(paras[5], C.TITLE)
set_para_text(paras[8], "Artificial Intelligence and Applications")
set_para_text(paras[9], "2026, Vol. XX(XX) 1-12")
set_para_text(paras[10], "DOI: 10.47852/bonviewAIAXXXXXXXX")

def set_author_line(p, authors):
    # Clear existing runs, then rebuild with superscripted affiliation numbers
    # and an asterisk only on the actual corresponding author.
    for r in list(p.runs):
        r.text = ""
    n = len(authors)
    for i, a in enumerate(authors):
        if i == 0:
            sep = ""
        elif i == n - 1:
            sep = ", and " if n > 2 else " and "
        else:
            sep = ", "
        p.add_run(sep + a["name"])
        sup = p.add_run("".join(str(x) for x in a["affil_idx"]))
        sup.font.superscript = True
        if a.get("corresponding"):
            star = p.add_run("*")
            star.font.superscript = True


set_author_line(paras[14], C.AUTHORS)

# Affiliation paragraphs 15,16,17 -> we have 3 affiliations; reuse those 3 paragraphs.
for i, aff in enumerate(C.AFFILIATIONS):
    set_para_text(paras[15 + i], f"{i + 1} {aff}.")

set_para_text(paras[18], C.CORRESPONDING_AUTHOR_NOTE)

set_para_text(paras[21], "Abstract: " + C.ABSTRACT)
set_para_text(paras[23], "Keywords: " + C.KEYWORDS)

# ---------------------------------------------------------------------------
# 2. Delete old body paragraphs (idx 25..end) and old sample tables.
# ---------------------------------------------------------------------------
body = doc.element.body
KEEP_UP_TO = paras[24]._p  # last paragraph we keep (blank line w/ the horizontal-rule drawing + sectPr boundary)

remove = False
for child in list(body):
    if child is KEEP_UP_TO:
        remove = True
        continue
    if remove:
        if child.tag == qn('w:sectPr'):
            # this is the doc-level trailing sectPr; keep it (it will now trail our new content)
            continue
        body.remove(child)

print("Body cleared of sample content.", file=sys.stderr)


# ---------------------------------------------------------------------------
# Helpers for building new content
# ---------------------------------------------------------------------------

def add_para(text, style=None, align=None, space_after=None):
    p = doc.add_paragraph(style=style)
    p.add_run(resolve_citations(text))
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(text, level):
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    return add_para(text, style=style)


def add_bullets(items):
    for head, body_text in items:
        p = doc.add_paragraph(style="List Bullet")
        r1 = p.add_run(head + " ")
        r1.bold = True
        p.add_run(resolve_citations(body_text))


def set_columns(section, num):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = sectPr.makeelement(qn('w:cols'), {})
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num))
    if num == 1:
        if qn('w:equalWidth') in cols.attrib:
            del cols.attrib[qn('w:equalWidth')]
    else:
        cols.set(qn('w:space'), "708")


def switch_columns(num):
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(new_section, num)
    return new_section


def add_full_width_block(build_fn):
    """Wrap build_fn() (which adds paragraphs/tables/pictures) in a 1-col region,
    then switch back to 2-col afterward."""
    switch_columns(1)
    build_fn()
    switch_columns(2)


def add_figure(path, caption, full):
    import os
    fp = os.path.join(FIGDIR, path)
    width = Inches(6.6) if full else Inches(3.15)

    def _build():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fp, width=width)
        cap = add_para(caption, style="FigureCaption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if full:
        add_full_width_block(_build)
    else:
        _build()


def add_code(lines, caption, full):
    def _build():
        for ln in lines:
            p = doc.add_paragraph(style="DispCompCode")
            r = p.add_run(ln if ln.strip() else " ")
            r.bold = False
            r.font.size = Pt(8.5)
        if caption:
            add_para(caption, style="FigureCaption")

    if full:
        add_full_width_block(_build)
    else:
        _build()


def add_algo(block):
    def _build():
        add_para(block["title"], style="AlgorithmTitle")
        p = doc.add_paragraph(style="Algorithm")
        r = p.add_run("Require: ")
        r.bold = True
        p.add_run(resolve_citations(block["require"]))
        p2 = doc.add_paragraph(style="Algorithm")
        r2 = p2.add_run("Ensure: ")
        r2.bold = True
        p2.add_run(resolve_citations(block["ensure"]))
        for i, step in enumerate(block["steps"], 1):
            sp = doc.add_paragraph(style="Algorithm")
            sp.add_run(f"{i}: {step}")
        rp = doc.add_paragraph(style="Algorithm")
        rr = rp.add_run("Return: ")
        rr.bold = True
        rp.add_run(block["ret"])

    add_full_width_block(_build)


def set_cell_text(cell, text, bold=False, size=8.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(resolve_citations(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    if align is not None:
        p.alignment = align


def add_table(block):
    header = block["header"]
    rows = block["rows"]
    note = block.get("note")

    def _build():
        add_para(block["caption"], style="TableTitle")
        t = doc.add_table(rows=1 + len(rows), cols=len(header))
        t.style = doc.styles["Table Grid"]
        t.autofit = True
        for j, h in enumerate(header):
            set_cell_text(t.cell(0, j), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        for i, row in enumerate(rows, start=1):
            is_last = (i == len(rows))
            bold_row = is_last and note == "bold_last_row"
            for j, val in enumerate(row):
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell_text(t.cell(i, j), val, bold=bold_row, align=align)
        if block.get("colw"):
            total_w = Inches(6.6)
            for j, frac in enumerate(block["colw"]):
                for row_cells in t.columns[j].cells:
                    row_cells.width = Inches(6.6 * frac)

    if block.get("full", True):
        add_full_width_block(_build)
    else:
        _build()


# ---------------------------------------------------------------------------
# 3. Emit BLOCKS
# ---------------------------------------------------------------------------
for b in C.BLOCKS:
    t = b["type"]
    if t == "appendix_start":
        pass  # no-op for DOCX; heading text already carries an explicit "Appendix X." prefix
    elif t == "h1":
        add_heading(resolve_citations(b["text"]), 1)
    elif t == "h2":
        add_heading(resolve_citations(b["text"]), 2)
    elif t == "h3":
        add_heading(resolve_citations(b["text"]), 3)
    elif t == "p":
        add_para(b["text"])
    elif t == "bullets":
        add_bullets(b["items"])
    elif t == "figure":
        add_figure(b["path"], resolve_citations(b["caption"]), b["full"])
    elif t == "code":
        add_code(b["lines"], resolve_citations(b["caption"]) if b.get("caption") else None, b["full"])
    elif t == "algo":
        add_algo(b)
    elif t == "table":
        add_table(b)
    else:
        raise ValueError(t)

# ---------------------------------------------------------------------------
# 4. References section
# ---------------------------------------------------------------------------
add_heading("References", 1)
for k in order:
    n = key_to_num[k]
    p = doc.add_paragraph(style="references")
    p.add_run(f"[{n}] {CITMAP[k]}")

# ---------------------------------------------------------------------------
import os
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("Saved:", OUT, file=sys.stderr)
print("Total citations:", len(order), file=sys.stderr)
